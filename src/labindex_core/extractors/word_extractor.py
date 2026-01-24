"""
Word document extractor.

Handles .docx files using python-docx.
"""

from pathlib import Path
from typing import List

from .base import TextExtractor, ExtractionResult


class WordExtractor(TextExtractor):
    """Extract text from Word documents."""

    EXTENSIONS = ['.docx', '.doc']

    def _extract_impl(self, path: Path) -> ExtractionResult:
        """Extract text from Word document."""
        ext = path.suffix.lower()

        if ext == '.docx':
            return self._extract_docx(path)
        elif ext == '.doc':
            # .doc format requires different library (antiword or similar)
            return ExtractionResult.failure(".doc format not supported (use .docx)")
        else:
            return ExtractionResult.failure(f"Unknown Word format: {ext}")

    def _extract_docx(self, path: Path) -> ExtractionResult:
        """Extract from .docx using python-docx."""
        try:
            from docx import Document
        except ImportError:
            return ExtractionResult.failure("python-docx not installed")

        try:
            doc = Document(path)

            paragraphs: List[str] = []

            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Extract text from tables
            table_texts: List[str] = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_texts.append(' | '.join(row_text))

            # Combine all text
            all_text = '\n'.join(paragraphs)
            if table_texts:
                all_text += '\n\n[Tables]\n' + '\n'.join(table_texts)

            return ExtractionResult(
                text=all_text,
                metadata={
                    'paragraph_count': len(paragraphs),
                    'table_count': len(doc.tables)
                },
                sources={'body': all_text[:500]}
            )

        except Exception as e:
            return ExtractionResult.failure(f"Word read error: {e}")
